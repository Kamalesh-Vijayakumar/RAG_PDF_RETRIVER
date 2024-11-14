from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

from docx import Document as DocxDocument
import PyPDF2
from typing import Union

# Import the existing functionalities
from llama_index.core import Document as Document, VectorStoreIndex
from llama_index.llms.mistralai import MistralAI
from llama_index.core import Settings

app = FastAPI()

# CORS middleware to allow connections from your React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Adjust in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

from io import BytesIO

def read_pdf(file_content):
    file_stream = BytesIO(file_content)
    reader = PyPDF2.PdfReader(file_stream)
    text = [page.extract_text() for page in reader.pages if page.extract_text() is not None]
    return "\n\n".join(text)

def read_docx(file_content):
    file_stream = BytesIO(file_content)
    doc = DocxDocument(file_stream)
    full_text = [para.text for para in doc.paragraphs if para.text]
    return "\n\n".join(full_text)


@app.get("/")
def read_root():
    return {"Hello": "World"}


from fastapi import HTTPException
import traceback

@app.post("/analyze/")
async def create_upload_file(file: UploadFile = File(...)):
    try:
        if not file.filename.lower().endswith(('.pdf', '.docx')):
            raise HTTPException(status_code=400, detail="Unsupported file type.")

        file_content = await file.read()
        if file.filename.lower().endswith('.pdf'):
            document_text = read_pdf(file_content)
        elif file.filename.lower().endswith('.docx'):
            document_text = read_docx(file_content)

        document = Document(text=document_text)
        settings = Settings
        settings.llm = MistralAI(api_key="IHCTcGVMsMqESQKsjDbQIEWt4v8hvIA6", model="mistral-large-latest")
        settings.embed_model = "local:BAAI/bge-small-en-v1.5"

        index = VectorStoreIndex.from_documents([document], settings=settings)
        query_engine = index.as_query_engine()
        response = query_engine.query("what is the document about?")

        return {"response": str(response)}

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
    
    
'''from fastapi import FastAPI, UploadFile, File, HTTPException, Body
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
from io import BytesIO
import os
import traceback
from docx import Document as DocxDocument
import PyPDF2

from llama_index.core import Document as LlamaDocument, VectorStoreIndex
from llama_index.llms.mistralai import MistralAI
from llama_index.core import Settings

app = FastAPI()

# Configure CORS middlewarefrom llama_index.core import Settings
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Adjust as needed for security
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory to temporarily store uploaded files
temp_storage_path = "./temp_storage"
os.makedirs(temp_storage_path, exist_ok=True)

# Global settings
llm_model = MistralAI(api_key="IHCTcGVMsMqESQKsjDbQIEWt4v8hvIA6", model="mistral-large-latest")
embed_model = "local:BAAI/bge-small-en-v1.5"  # Embed model can be defined globally

def read_pdf(file_content: bytes) -> str:
    file_stream = BytesIO(file_content)
    reader = PyPDF2.PdfReader(file_stream)
    text = [page.extract_text() for page in reader.pages if page.extract_text() is not None]
    return "\n\n".join(text)

def read_docx(file_content: bytes) -> str:
    file_stream = BytesIO(file_content)
    doc = DocxDocument(file_stream)
    full_text = [para.text for para in doc.paragraphs if para.text]
    return "\n\n".join(full_text)

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    try:
        file_extension = file.filename.split('.')[-1].lower()
        if file_extension not in ['pdf', 'docx']:
            raise HTTPException(status_code=400, detail="Unsupported file type. Only PDF and DOCX are allowed.")
        
        temp_file_path = os.path.join(temp_storage_path, file.filename)
        with open(temp_file_path, 'wb') as f:
            f.write(await file.read())
        
        return {"filename": temp_file_path}
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze/")
async def analyze_document(data: dict = Body(...)):
    query = data.get('query')
    filename = data.get('filename')
    
    if not query or not filename:
        raise HTTPException(status_code=400, detail="Both 'query' and 'filename' are required.")
    
    try:
        with open(filename, "rb") as file:
            file_content = file.read()
            if filename.endswith('.pdf'):
                document_text = read_pdf(file_content)
            elif filename.endswith('.docx'):
                document_text = read_docx(file_content)
            else:
                raise HTTPException(status_code=400, detail="Unsupported file type.")
        
        # Initialize the LlamaIndex document and index
        document = LlamaDocument(text=document_text)
        index = VectorStoreIndex.from_documents([document])
        
        # Use the embed model directly without Pydantic
        Settings.embed_model = embed_model  # If needed for LlamaIndex usage
        query_engine = index.as_query_engine()  # Get the query engine
        
        # Execute the query
        response = query_engine.query(query)
        
        return {"response": response}
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
'''