from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
from io import BytesIO
import os
import traceback
from docx import Document as DocxDocument
import PyPDF2
from llama_index.core import Document as LlamaDocument, VectorStoreIndex
from llama_index.llms.mistralai import MistralAI
from llama_index.core import Settings
from pydantic import BaseModel, Field , validator

# Secret management with Pydantic v2
class MyModel(BaseModel):
    secret: str = Field(..., repr=False)
    
    class Config:
        json_encoders = {str: lambda v: "[SECRET]"}  # Hides the secret value

    def get_secret_value(self):
        return self.secret  # Accessing the secret value

# FastAPI app
app = FastAPI()

# Configure CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Adjust as needed for security
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory to temporarily store uploaded files
temp_storage_path = "./temp_storage"
os.makedirs(temp_storage_path, exist_ok=True)

# Global settings for language and embedding models
llm_model = MistralAI(api_key="API_KEY", model="mistral-large-latest")
embed_model = "local:BAAI/bge-small-en-v1.5"  # Embed model can be defined globally

def read_pdf(file_content: bytes) -> str:
    """Read and extract text from a PDF file."""
    file_stream = BytesIO(file_content)
    reader = PyPDF2.PdfReader(file_stream)
    text = [page.extract_text() for page in reader.pages if page.extract_text() is not None]
    return "\n\n".join(text)

def read_docx(file_content: bytes) -> str:
    """Read and extract text from a DOCX file."""
    file_stream = BytesIO(file_content)
    doc = DocxDocument(file_stream)
    full_text = [para.text for para in doc.paragraphs if para.text]
    return "\n\n".join(full_text)

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    """Endpoint to handle file uploads."""
    try:
        file_extension = file.filename.split('.')[-1].lower()
        if file_extension not in ['pdf', 'docx']:
            raise HTTPException(status_code=400, detail="Unsupported file type. Only PDF and DOCX are allowed.")
        
        # Save file to temp storage
        temp_file_path = os.path.join(temp_storage_path, file.filename)
        with open(temp_file_path, 'wb') as f:
            f.write(await file.read())
        
        return {"filename": temp_file_path}
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

class AnalyzeRequest(BaseModel):
    query: str
    filename: str

    @validator('filename')
    def validate_filename(cls, v):
        if not os.path.exists(v):
            raise ValueError("File not found.")
        return v

@app.post("/analyze/")
async def analyze_document(data: AnalyzeRequest):
    """Endpoint to analyze a document with a specified query."""
    query = data.query
    filename = data.filename
    
    try:
        # Read file content based on file type
        with open(filename, "rb") as file:
            file_content = file.read()
            if filename.endswith('.pdf'):
                document_text = read_pdf(file_content)
            elif filename.endswith('.docx'):
                document_text = read_docx(file_content)
            else:
                raise HTTPException(status_code=400, detail="Unsupported file type.")
        
        # Initialize the LlamaIndex document and index
        document = LlamaDocument(text=document_text)
        index = VectorStoreIndex.from_documents([document])
        
        # Set the embedding model directly without Pydantic validation issues
        Settings.embed_model = embed_model
        query_engine = index.as_query_engine()  # Get the query engine
        
        # Execute the query
        response = query_engine.query(query)
        
        return {"response": response}
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)

    # Example usage of MyModel
    model_instance = MyModel(secret="my_secret")
    print(model_instance.get_secret_value())  # Accessing the secret value
